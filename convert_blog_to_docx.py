#!/usr/bin/env python3
"""
Convert BLOG_POST_AGENTIC_AI_ADOPTION.md to a richly formatted Microsoft Word (.docx).
Preserves all content, formatting, structure, and visual hierarchy with professional styling.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# File paths
md_file = Path("/Users/arun.ray/personal-projects/a1-agent-engine/docs/blogs/BLOG_POST_AGENTIC_AI_ADOPTION.md")
docx_file = Path("/Users/arun.ray/personal-projects/a1-agent-engine/docs/blogs/BLOG_POST_AGENTIC_AI_ADOPTION.docx")

# Read markdown
content = md_file.read_text(encoding='utf-8')
lines = content.split('\n')

# Create document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

def add_shading(element, color_hex):
    """Add background shading to a paragraph or run."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    element._element.get_or_add_pPr().append(shading_elm)

def add_borders(para, color_hex='E0E0E0', width='4'):
    """Add borders around a paragraph."""
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), width)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color_hex)
        pBdr.append(border)
    pPr.append(pBdr)

def add_formatted_text(para, text):
    """Add text with inline markdown formatting to a paragraph."""
    if not text:
        return

    # Pattern for: **bold**, *italic*, `code`, or regular text
    pattern = r'(\*\*[^\*]+\*\*|\*[^\*]+\*|`[^`]+`|[^*`]+)'

    for match in re.finditer(pattern, text):
        segment = match.group(0)

        if segment.startswith('**') and segment.endswith('**'):
            # Bold
            run = para.add_run(segment[2:-2])
            run.bold = True
            run.font.color.rgb = RGBColor(26, 26, 26)
        elif segment.startswith('*') and segment.endswith('*'):
            # Italic
            run = para.add_run(segment[1:-1])
            run.italic = True
            run.font.color.rgb = RGBColor(26, 26, 26)
        elif segment.startswith('`') and segment.endswith('`'):
            # Code
            run = para.add_run(segment[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(44, 62, 80)
        else:
            # Normal text
            run = para.add_run(segment)
            run.font.color.rgb = RGBColor(26, 26, 26)

# Add title
title = doc.add_paragraph()
title_run = title.add_run("Agentic AI in Regulated Enterprises")
title_run.font.size = Pt(26)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(26, 26, 26)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(6)

# Add subtitle
subtitle = doc.add_paragraph("The Case for Sovereign, Compliant, Verifiable Automation")
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = RGBColor(60, 62, 80)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(24)

# Parse markdown and add to document
i = 0

while i < len(lines):
    line = lines[i]

    # Skip empty lines
    if not line.strip():
        i += 1
        continue

    # Heading 1 (# but not ##)
    if line.startswith('# ') and not line.startswith('## '):
        heading_text = line[2:].strip()
        heading = doc.add_heading(heading_text, level=1)
        heading.paragraph_format.space_after = Pt(12)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(26, 26, 26)
        i += 1
        continue

    # Heading 2 (##)
    if line.startswith('## '):
        heading_text = line[3:].strip()
        heading = doc.add_heading(heading_text, level=2)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(8)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(44, 62, 80)
        i += 1
        continue

    # Heading 3 (###)
    if line.startswith('### '):
        heading_text = line[4:].strip()
        heading = doc.add_heading(heading_text, level=3)
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(6)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(52, 73, 94)
        i += 1
        continue

    # Heading 4 (####)
    if line.startswith('#### '):
        heading_text = line[5:].strip()
        heading = doc.add_heading(heading_text, level=4)
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(4)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(52, 73, 94)
        i += 1
        continue

    # Horizontal rule (---)
    if line.strip() == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)
        i += 1
        continue

    # Code block (triple backticks)
    if line.strip().startswith('```'):
        code_lines = []
        i += 1

        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        i += 1  # skip closing ```

        if code_lines:
            # Add code block with formatting
            code_para = doc.add_paragraph()
            code_para.style = 'Normal'
            code_para.paragraph_format.left_indent = Inches(0.3)
            code_para.paragraph_format.space_after = Pt(8)
            add_shading(code_para, 'F8F8F8')
            add_borders(code_para, 'D0D0D0', '8')

            # Add code text
            for code_line in code_lines:
                if code_line or code_lines[-1] != code_line:  # Don't add trailing empty lines
                    run = code_para.add_run(code_line + '\n')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(44, 62, 80)

        continue

    # Unordered list (-)
    if line.strip().startswith('- '):
        bullet_text = line.strip()[2:].strip()
        list_para = doc.add_paragraph(style='List Bullet')
        list_para.paragraph_format.space_after = Pt(2)
        add_formatted_text(list_para, bullet_text)
        i += 1
        continue

    # Ordered list (number.)
    if re.match(r'^\s*\d+\. ', line):
        match = re.match(r'^\s*(\d+)\. (.*)', line)
        if match:
            list_text = match.group(2).strip()
            list_para = doc.add_paragraph(style='List Number')
            list_para.paragraph_format.space_after = Pt(2)
            add_formatted_text(list_para, list_text)
        i += 1
        continue

    # Table (starts with |)
    if line.strip().startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1

        # Parse table rows
        rows = []
        for row in table_lines:
            cells = [cell.strip() for cell in row.split('|')[1:-1]]
            if cells:
                rows.append(cells)

        # Skip separator row (contains dashes)
        if len(rows) > 1 and all(re.match(r'^\-+$', cell) for cell in rows[1]):
            rows = [rows[0]] + rows[2:]

        if rows:
            # Create table
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Light Grid Accent 1'

            for row_idx, row_cells in enumerate(rows):
                for col_idx, cell_text in enumerate(row_cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell_para = cell.paragraphs[0]
                    cell_para.paragraph_format.space_after = Pt(4)

                    # Header row
                    if row_idx == 0:
                        cell_para.clear()
                        run = cell_para.add_run(cell_text)
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), '4472C4')
                        cell._element.get_or_add_tcPr().append(shading_elm)
                    else:
                        cell_para.clear()
                        add_formatted_text(cell_para, cell_text)

            # Add space after table
            doc.add_paragraph()

        continue

    # Regular paragraph
    if line.strip():
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.15
        add_formatted_text(para, line.strip())

    i += 1

# Save document
doc.save(str(docx_file))
print(f"✅ Word document (.docx) created: {docx_file}")
print(f"📄 File size: {docx_file.stat().st_size / 1024 / 1024:.2f} MB")
print(f"📝 Source markdown: {md_file}")
print(f"📑 Professional formatting with styles, colors, and rich text applied")
